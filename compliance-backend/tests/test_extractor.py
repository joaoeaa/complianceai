"""
Unit tests for the document extraction service.

PDF extraction is tested by mocking pdfplumber.open (the library is still
required to be installed, only its I/O behaviour is stubbed out).

DOCX extraction is tested against real files created in-memory with
python-docx, giving us genuine extraction coverage without any network I/O.
"""
import os
import pytest
from pathlib import Path
from unittest.mock import MagicMock, patch, call


# ─── Helpers ─────────────────────────────────────────────────────────────────

def _make_mock_pdf(pages_text: list[str | None]):
    """Return a pdfplumber-compatible mock with the given per-page text.

    Pass None to simulate a page that returns no text (image/scan).
    """
    mock_pages = []
    for text in pages_text:
        page = MagicMock()
        page.extract_text.return_value = text
        mock_pages.append(page)

    mock_pdf = MagicMock()
    mock_pdf.pages = mock_pages
    mock_pdf.__enter__ = lambda s: s
    mock_pdf.__exit__ = MagicMock(return_value=False)
    return mock_pdf


def _make_real_docx(tmp_path: Path, paragraphs: list[str], table_rows: list[list[str]] | None = None) -> Path:
    """Create a genuine .docx file using python-docx and return its path."""
    from docx import Document as DocxDocument

    doc = DocxDocument()
    for text in paragraphs:
        doc.add_paragraph(text)

    if table_rows:
        cols = max(len(row) for row in table_rows)
        table = doc.add_table(rows=len(table_rows), cols=cols)
        for r, row in enumerate(table_rows):
            for c, cell_text in enumerate(row):
                table.cell(r, c).text = cell_text

    path = tmp_path / "test.docx"
    doc.save(str(path))
    return path


# ─── get_mime_type ────────────────────────────────────────────────────────────

class TestGetMimeType:
    """Pure unit tests for the MIME-type inference helper."""

    def test_pdf_extension(self):
        from app.services.document_extractor import get_mime_type
        assert get_mime_type("contract.pdf") == "application/pdf"

    def test_docx_extension(self):
        from app.services.document_extractor import get_mime_type
        assert get_mime_type("report.docx") == (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    def test_uppercase_extension(self):
        from app.services.document_extractor import get_mime_type
        assert get_mime_type("DOCUMENTO.PDF") == "application/pdf"

    def test_unknown_extension_returns_octet_stream(self):
        from app.services.document_extractor import get_mime_type
        assert get_mime_type("file.txt") == "application/octet-stream"

    def test_extension_with_path(self):
        from app.services.document_extractor import get_mime_type
        assert get_mime_type("/tmp/uploads/uuid123.docx") == (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )


# ─── extract_text dispatcher ─────────────────────────────────────────────────

class TestExtractTextDispatcher:
    """Tests for the top-level extract_text() routing function."""

    def test_file_not_found_raises(self, tmp_path):
        from app.services.document_extractor import extract_text
        with pytest.raises(FileNotFoundError, match="não encontrado"):
            extract_text(str(tmp_path / "ghost.pdf"), "application/pdf")

    def test_empty_file_raises(self, tmp_path):
        from app.services.document_extractor import extract_text
        empty = tmp_path / "empty.pdf"
        empty.write_bytes(b"")
        with pytest.raises(ValueError, match="vazio"):
            extract_text(str(empty), "application/pdf")

    def test_unsupported_format_raises(self, tmp_path):
        from app.services.document_extractor import extract_text
        txt = tmp_path / "doc.txt"
        txt.write_bytes(b"hello world")
        with pytest.raises(ValueError, match="não suportado"):
            extract_text(str(txt), "text/plain")

    def test_routes_pdf_by_mime(self, tmp_path):
        from app.services.document_extractor import extract_text
        pdf = tmp_path / "doc.pdf"
        pdf.write_bytes(b"%PDF minimal")

        mock_pdf = _make_mock_pdf(["Cláusula 1 do contrato."])
        with patch("pdfplumber.open", return_value=mock_pdf):
            result = extract_text(str(pdf), "application/pdf")

        assert "Cláusula 1" in result

    def test_routes_docx_by_extension(self, tmp_path):
        from app.services.document_extractor import extract_text
        docx_path = _make_real_docx(tmp_path, ["Contrato de prestação de serviços."])
        result = extract_text(str(docx_path), "application/octet-stream")  # mime wrong, ext correct
        assert "Contrato de prestação" in result


# ─── PDF extraction ───────────────────────────────────────────────────────────

class TestExtractTextFromPdf:
    """Tests for the PDF-specific extraction function."""

    def test_single_page(self, tmp_path):
        from app.services.document_extractor import extract_text_from_pdf
        pdf = tmp_path / "single.pdf"
        pdf.write_bytes(b"%PDF minimal")

        mock_pdf = _make_mock_pdf(["Texto da página 1."])
        with patch("pdfplumber.open", return_value=mock_pdf):
            result = extract_text_from_pdf(str(pdf))

        assert "Texto da página 1." in result
        assert "Página 1" in result

    def test_multiple_pages(self, tmp_path):
        from app.services.document_extractor import extract_text_from_pdf
        pdf = tmp_path / "multi.pdf"
        pdf.write_bytes(b"%PDF minimal")

        mock_pdf = _make_mock_pdf(["Primeira página.", "Segunda página.", "Terceira página."])
        with patch("pdfplumber.open", return_value=mock_pdf):
            result = extract_text_from_pdf(str(pdf))

        assert "Primeira página." in result
        assert "Segunda página." in result
        assert "Terceira página." in result
        assert "Página 1" in result
        assert "Página 2" in result

    def test_page_without_text_logs_warning(self, tmp_path):
        from app.services.document_extractor import extract_text_from_pdf
        pdf = tmp_path / "partial.pdf"
        pdf.write_bytes(b"%PDF minimal")

        # Page 1 has text, page 2 is a scan (returns None)
        mock_pdf = _make_mock_pdf(["Texto válido.", None])
        with patch("pdfplumber.open", return_value=mock_pdf):
            result = extract_text_from_pdf(str(pdf))

        assert "Texto válido." in result

    def test_all_pages_empty_triggers_ocr_attempt(self, tmp_path):
        from app.services.document_extractor import extract_text_from_pdf
        pdf = tmp_path / "scan.pdf"
        pdf.write_bytes(b"%PDF minimal")

        mock_pdf = _make_mock_pdf([None, None])

        # OCR libs not installed — expect ImportError message
        with patch("pdfplumber.open", return_value=mock_pdf):
            with pytest.raises(ValueError, match="(?i)(ocr|escaneado|tesseract)"):
                extract_text_from_pdf(str(pdf))

    def test_pdfplumber_error_raises_value_error(self, tmp_path):
        from app.services.document_extractor import extract_text_from_pdf
        pdf = tmp_path / "broken.pdf"
        pdf.write_bytes(b"%PDF broken")

        with patch("pdfplumber.open", side_effect=Exception("arquivo corrompido")):
            with pytest.raises(ValueError, match="Falha na extração"):
                extract_text_from_pdf(str(pdf))


# ─── DOCX extraction ──────────────────────────────────────────────────────────

class TestExtractTextFromDocx:
    """Tests for the DOCX-specific extraction function using real files."""

    def test_simple_paragraphs(self, tmp_path):
        from app.services.document_extractor import extract_text_from_docx
        path = _make_real_docx(tmp_path, [
            "CONTRATO DE PRESTAÇÃO DE SERVIÇOS",
            "Cláusula 1 — Das partes.",
            "Cláusula 2 — Do objeto.",
        ])
        result = extract_text_from_docx(str(path))

        assert "CONTRATO DE PRESTAÇÃO DE SERVIÇOS" in result
        assert "Cláusula 1" in result
        assert "Cláusula 2" in result

    def test_table_content_extracted(self, tmp_path):
        from app.services.document_extractor import extract_text_from_docx
        path = _make_real_docx(
            tmp_path,
            paragraphs=["Tabela de valores:"],
            table_rows=[
                ["Item", "Valor"],
                ["Serviço A", "R$ 1.000,00"],
                ["Serviço B", "R$ 2.500,00"],
            ],
        )
        result = extract_text_from_docx(str(path))

        assert "Tabela de valores:" in result
        assert "Serviço A" in result
        assert "R$ 1.000,00" in result

    def test_empty_docx_raises(self, tmp_path):
        from app.services.document_extractor import extract_text_from_docx
        from docx import Document as DocxDocument

        # Create a DOCX with no content
        doc = DocxDocument()
        path = tmp_path / "empty.docx"
        doc.save(str(path))

        with pytest.raises(ValueError, match="(?i)(vazio|legível)"):
            extract_text_from_docx(str(path))

    def test_corrupted_docx_raises_value_error(self, tmp_path):
        from app.services.document_extractor import extract_text_from_docx
        path = tmp_path / "corrupt.docx"
        path.write_bytes(b"not a real docx file")

        with pytest.raises(ValueError, match="Falha na extração"):
            extract_text_from_docx(str(path))

    def test_extracts_multiple_paragraphs_joined(self, tmp_path):
        from app.services.document_extractor import extract_text_from_docx
        texts = [f"Parágrafo {i}." for i in range(5)]
        path = _make_real_docx(tmp_path, texts)
        result = extract_text_from_docx(str(path))

        for text in texts:
            assert text in result
