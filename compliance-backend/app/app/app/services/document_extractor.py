"""
Document text extraction service.
Handles PDF (via pdfplumber) and DOCX (via python-docx).
"""
import os
import logging
from pathlib import Path
from typing import Tuple

logger = logging.getLogger(__name__)


def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from a PDF file using pdfplumber."""
    import pdfplumber

    text_parts = []
    try:
        with pdfplumber.open(file_path) as pdf:
            for i, page in enumerate(pdf.pages):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(f"--- Página {i + 1} ---\n{page_text}")
                else:
                    logger.warning(f"Página {i + 1} sem texto extraído (pode ser imagem/scan)")
    except Exception as e:
        logger.error(f"Erro ao extrair texto do PDF: {e}")
        raise ValueError(f"Falha na extração do PDF: {str(e)}")

    if not text_parts:
        # Fallback: attempt OCR (requires tesseract installed)
        logger.info("PDF sem texto detectado, tentando OCR...")
        return _ocr_fallback(file_path)

    return "\n\n".join(text_parts)


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from a DOCX file using python-docx."""
    from docx import Document

    text_parts = []
    try:
        doc = Document(file_path)
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)

    except Exception as e:
        logger.error(f"Erro ao extrair texto do DOCX: {e}")
        raise ValueError(f"Falha na extração do DOCX: {str(e)}")

    if not text_parts:
        raise ValueError("Documento DOCX está vazio ou não contém texto legível.")

    return "\n\n".join(text_parts)


def _ocr_fallback(file_path: str) -> str:
    """
    OCR fallback for scanned PDFs.
    Requires: tesseract, pdf2image, pillow
    """
    try:
        import pytesseract
        from pdf2image import convert_from_path

        images = convert_from_path(file_path, dpi=300)
        text_parts = []
        for i, img in enumerate(images):
            text = pytesseract.image_to_string(img, lang="por")  # Portuguese
            if text.strip():
                text_parts.append(f"--- Página {i + 1} (OCR) ---\n{text}")

        if not text_parts:
            raise ValueError("OCR não conseguiu extrair texto do documento.")

        return "\n\n".join(text_parts)

    except ImportError:
        raise ValueError(
            "PDF parece ser escaneado/imagem mas OCR (tesseract) não está instalado. "
            "Instale: apt-get install tesseract-ocr tesseract-ocr-por && pip install pdf2image pytesseract"
        )


def extract_text(file_path: str, mime_type: str) -> str:
    """
    Main extraction function. Routes to the correct extractor based on mime type.
    Returns the full extracted text.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Arquivo não encontrado: {file_path}")

    file_size = os.path.getsize(file_path)
    if file_size == 0:
        raise ValueError("Arquivo está vazio (0 bytes).")

    logger.info(f"Extraindo texto de {file_path} ({mime_type}, {file_size} bytes)")

    if mime_type == "application/pdf" or file_path.lower().endswith(".pdf"):
        return extract_text_from_pdf(file_path)
    elif (
        mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        or file_path.lower().endswith(".docx")
    ):
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Formato não suportado: {mime_type}")


def get_mime_type(filename: str) -> str:
    """Infer mime type from filename extension."""
    ext = Path(filename).suffix.lower()
    mapping = {
        ".pdf": "application/pdf",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }
    return mapping.get(ext, "application/octet-stream")
